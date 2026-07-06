#!/usr/bin/env python3
"""Generate Elastic Tree site content Word document for review."""

import json
import re
from pathlib import Path

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_LINE_SPACING

ROOT = Path(__file__).resolve().parents[1]
DATA = ROOT / "docs" / "content-export-data.json"
OUT = ROOT / "docs" / "Elastic-Tree-Site-Content.docx"


def style_doc(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    pf = normal.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = 1.15
    pf.space_after = Pt(6)


def add_title(doc: Document, text: str) -> None:
    doc.add_heading(text, 0)


def add_h1(doc: Document, text: str) -> None:
    doc.add_heading(text, 1)


def add_h2(doc: Document, text: str) -> None:
    doc.add_heading(text, 2)


def add_h3(doc: Document, text: str) -> None:
    doc.add_heading(text, 3)


def add_para(doc: Document, text: str) -> None:
    doc.add_paragraph(text)


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def strip_md(text: str) -> str:
    text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)
    text = re.sub(r"\*(.+?)\*", r"\1", text)
    return text


def add_markdown_content(doc: Document, content: str) -> None:
    for block in content.split("\n\n"):
        block = block.strip()
        if not block:
            continue
        if block.startswith("## "):
            add_h3(doc, strip_md(block[3:].strip()))
        elif block.startswith("### "):
            p = doc.add_paragraph()
            run = p.add_run(strip_md(block[4:].strip()))
            run.bold = True
        elif block.startswith("- "):
            lines = [strip_md(line[2:].strip()) for line in block.split("\n") if line.startswith("- ")]
            add_bullets(doc, lines)
        elif re.match(r"^\d+\.\s", block):
            for line in block.split("\n"):
                line = re.sub(r"^\d+\.\s+", "", line.strip())
                if line:
                    doc.add_paragraph(strip_md(line), style="List Number")
        else:
            add_para(doc, strip_md(block.replace("\n", " ")))


def add_section_fields(doc: Document, data: dict, prefix: str = "") -> None:
    for key, val in data.items():
        label = f"{prefix}{key.replace('_', ' ').title()}" if prefix else key.replace("_", " ").title()
        if isinstance(val, str):
            add_para(doc, f"{label}: {val}")
        elif isinstance(val, list):
            if val and isinstance(val[0], str):
                add_para(doc, f"{label}:")
                add_bullets(doc, val)
            elif val and isinstance(val[0], dict):
                add_h3(doc, label)
                for item in val:
                    if "title" in item and "desc" in item:
                        add_para(doc, f"{item['title']}: {item['desc']}")
                    elif "name" in item and "quote" in item:
                        add_para(doc, f"\"{item['quote']}\" — {item['name']}, {item.get('company', '')}")
                    elif "title" in item:
                        parts = [item.get("title", ""), item.get("tag", ""), item.get("desc", "")]
                        add_para(doc, " — ".join(p for p in parts if p))
                    else:
                        add_para(doc, json.dumps(item, ensure_ascii=False))
            else:
                add_para(doc, f"{label}: {', '.join(str(v) for v in val)}")
        elif isinstance(val, dict):
            add_h3(doc, label)
            add_section_fields(doc, val)


def main() -> None:
    data = json.loads(DATA.read_text(encoding="utf-8"))
    site = data["siteCopy"]
    doc = Document()
    style_doc(doc)

    add_title(doc, "Elastic Tree Website — Full Content Export")
    add_para(doc, "Review and edit this document. Mark changes or send an updated version back for implementation on the website.")
    add_para(doc, f"Site: {site['global']['siteUrl']} | Tagline: {site['global']['tagline']}")

    add_h1(doc, "Global Metadata")
    add_section_fields(doc, site["global"])

    add_h1(doc, "Navigation")
    add_bullets(doc, site["nav"])
    add_para(doc, "Navbar button: Get in Touch")
    add_para(doc, "Footer: Insight-led market research for FMCG and food-service brands. Smart decisions, simply made — since 2014.")
    add_para(doc, "Footer email: research@elastictree.com | HQ: Bengaluru, India")

    add_h1(doc, "Home Page")
    add_section_fields(doc, site["home"])

    add_h2(doc, "About stats")
    for s in data["aboutStats"]:
        add_para(doc, f"{s['value']} — {s['label']}: {s['desc']}")

    add_h2(doc, "Team members")
    for m in data["teamMembers"]:
        add_para(doc, f"{m['name']} — {m['role']}")

    add_h2(doc, "Offices")
    for o in data["offices"]:
        add_para(doc, f"{o['city']}: {o['address']}")

    add_h2(doc, "Client list (59)")
    add_para(doc, ", ".join(site["clients"]))

    add_h2(doc, "Testimonials")
    for t in site["testimonials"]:
        add_para(doc, f"\"{t['quote']}\"")
        add_para(doc, f"— {t['name']}, {t['company']}")
        doc.add_paragraph()

    add_h1(doc, "Capabilities Page")
    cap = site["capabilities"]
    add_section_fields(doc, cap["hero"])
    add_para(doc, "Sticky navigation: " + " | ".join(cap["stickyNav"]))
    for pillar in cap["pillars"]:
        add_h2(doc, pillar["label"])
        add_para(doc, f"Tagline: {pillar['tagline']}")
        add_para(doc, pillar["summary"])
        add_para(doc, "Used for:")
        add_bullets(doc, pillar["usedFor"])
        add_para(doc, "Methods & approaches:")
        add_bullets(doc, pillar["methods"])

    adv = cap["advancedMethods"]
    add_h2(doc, adv["title"])
    add_para(doc, adv["subtitle"])
    for tool in adv["tools"]:
        add_h3(doc, tool["title"])
        add_para(doc, tool["subtitle"])
        add_para(doc, tool["desc"])
        add_para(doc, tool["stats"])

    add_h3(doc, "AI Gaze™ — Four Deliverables")
    for d in adv["gazeDeliverables"]:
        add_para(doc, f"{d['name']}: {d['desc']}")

    add_h3(doc, "Where Clients Apply AI Gaze™")
    add_bullets(doc, adv["gazeUseCases"])

    add_h2(doc, "Capabilities CTA")
    add_section_fields(doc, cap["cta"])

    add_h1(doc, "Table Share Page")
    ts = site["tableSharePage"]
    add_section_fields(doc, ts["hero"])
    add_section_fields(doc, ts["decisionPlatform"])
    add_para(doc, "Deliverables:")
    add_bullets(doc, data["tableShareDeliverables"])
    add_section_fields(doc, ts["whyConsumption"])
    add_para(doc, "Insights:")
    for item in data["tableShareInsights"]:
        add_para(doc, f"{item['title']}: {item['desc']}")
    add_section_fields(doc, ts["occasions"])
    for occ in data["tableShareOccasions"]:
        add_para(doc, f"{occ['name']}: {occ['desc']}")
    add_section_fields(doc, ts["sourceOfPrep"])
    add_section_fields(doc, ts["volumeRelevance"])
    add_para(doc, "Business uses:")
    for use in data["tableShareBusinessUses"]:
        add_para(doc, f"{use['title']}: {use['desc']}")
    add_section_fields(doc, ts["geographicReach"])
    add_para(doc, "Cities: " + ", ".join(data["tableShareCities"]))
    add_h2(doc, "Table Share CTA")
    add_section_fields(doc, ts["cta"])

    add_h1(doc, "Contact Page")
    add_section_fields(doc, site["contact"])

    add_h1(doc, "Case Studies Page")
    csp = site["caseStudiesPage"]
    add_para(doc, f"Eyebrow: {csp['eyebrow']}")
    add_para(doc, f"Title: {csp['title']}")
    add_para(doc, f"Subtitle: {len(data['posts'])} published studies across analytics, sensory science, syndicated research, AI, and impact.")
    add_section_fields(doc, {k: v for k, v in csp.items() if k not in ("eyebrow", "title")})

    add_h1(doc, "All Case Studies & Articles")
    for i, post in enumerate(data["posts"], 1):
        add_h2(doc, f"{i}. {post['title']}")
        meta = [f"Tags: {', '.join(post['tags'])}", f"Read time: {post['readTime']}", f"Date: {post['date']}"]
        if post.get("client"):
            meta.append(f"Client: {post['client']}")
        if post.get("industry"):
            meta.append(f"Industry: {post['industry']}")
        add_para(doc, " | ".join(meta))
        add_para(doc, f"Excerpt: {post['excerpt']}")
        add_h3(doc, "Full content")
        add_markdown_content(doc, post["content"])
        doc.add_paragraph()

    doc.save(OUT)
    print(f"Wrote {OUT}")


if __name__ == "__main__":
    main()
